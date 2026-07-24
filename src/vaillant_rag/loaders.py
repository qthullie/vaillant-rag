"""Text extraction from source documents.

Dispatches on file extension. Supported formats:

- ``.pdf`` — via PyMuPDF, with optional OCR of embedded images
  (requires the ``[ocr]`` extra and a Tesseract installation).
- ``.txt``, ``.md``, ``.markdown``, ``.rst`` — read as plain text.
- ``.html``, ``.htm`` — tags stripped with the standard-library parser.
- ``.docx`` — via python-docx (requires the ``[office]`` extra);
  headings and tables are rendered in Markdown flavor.
- ``.xlsx``, ``.xlsm`` — via openpyxl (requires the ``[office]`` extra);
  one Markdown table per sheet.

Legacy binary formats (``.doc``, ``.xls``) are not supported; convert
them to ``.docx`` / ``.xlsx`` first.

:func:`extract_text` is the single entry point, producing flat text for
the embedding pipeline. To support a new format, add an entry to
:data:`_LOADERS`.
"""

from __future__ import annotations

import io
import logging
import re
from collections.abc import Callable
from html.parser import HTMLParser
from pathlib import Path

logger = logging.getLogger(__name__)

# Optional OCR dependencies -- absence only disables OCR, never extraction.
try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]

# OCR pre-processing constants.
_OCR_UPSCALE_FACTOR = 1.5
_OCR_BINARIZE_THRESHOLD = 180

_OFFICE_EXTRA_HINT = "pip install vaillant-rag[office]"


def _ocr_image(img: Image.Image) -> str:
    """Run OCR on a PIL image with light pre-processing.

    Returns an empty string when OCR is unavailable or fails; OCR is a
    best-effort enrichment and must never break extraction.
    """
    if not (pytesseract and Image):
        return ""
    try:
        gray = img.convert("L")
        width, height = gray.size
        gray = gray.resize((int(width * _OCR_UPSCALE_FACTOR), int(height * _OCR_UPSCALE_FACTOR)))
        binary = gray.point(lambda x: 255 if x > _OCR_BINARIZE_THRESHOLD else 0, mode="1")
        return pytesseract.image_to_string(binary, config="--oem 3 --psm 6")
    except Exception as exc:
        logger.debug("OCR failed on embedded image: %s", exc)
        return ""


def _extract_pdf(path: Path, ocr_images: bool = False) -> str:
    """Extract text from a PDF, optionally OCR-ing embedded images."""
    import fitz  # PyMuPDF; imported lazily to keep non-PDF usage light

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page_number, page in enumerate(doc, start=1):
            parts.append(page.get_text("text"))
            if not (ocr_images and pytesseract and Image):
                continue
            for image_info in page.get_images(full=True):
                xref = image_info[0]
                try:
                    pixmap = fitz.Pixmap(doc, xref)
                    if pixmap.n >= 4:  # convert CMYK/alpha to RGB
                        pixmap = fitz.Pixmap(fitz.csRGB, pixmap)
                    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                    ocr_text = _ocr_image(image).strip()
                    if ocr_text:
                        parts.append(f"[Figure p.{page_number}] {ocr_text}")
                except Exception as exc:
                    logger.warning(
                        "Skipping unreadable image (xref=%d, page=%d) in %s: %s",
                        xref,
                        page_number,
                        path.name,
                        exc,
                    )
    return "\n".join(parts)


def _extract_plain_text(path: Path, ocr_images: bool = False) -> str:
    """Read a plain-text file (UTF-8, tolerant of stray bytes)."""
    return path.read_text(encoding="utf-8", errors="replace")


class _HTMLTextExtractor(HTMLParser):
    """Collect text nodes, skipping script/style contents."""

    _SKIPPED_TAGS = {"script", "style"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._SKIPPED_TAGS:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIPPED_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth > 0 or not data.strip():
            return
        self._parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(self._parts)


def _extract_html(path: Path, ocr_images: bool = False) -> str:
    """Extract visible text from an HTML file."""
    parser = _HTMLTextExtractor()
    parser.feed(path.read_text(encoding="utf-8", errors="replace"))
    return parser.text()


# Word heading styles carry localized names ("Heading 1", "Titre 1", ...);
# match on the trailing level digit in either the name or the style id.
_DOCX_HEADING_STYLE = re.compile(r"(?:heading|titre|berschrift|t[ií]tulo)\s*(\d)", re.IGNORECASE)


def _docx_heading_level(paragraph: object) -> int | None:
    """Heading level of a python-docx paragraph, or None for body text."""
    style = getattr(paragraph, "style", None)
    if style is None:
        return None
    for candidate in (getattr(style, "name", "") or "", getattr(style, "style_id", "") or ""):
        match = _DOCX_HEADING_STYLE.search(candidate)
        if match:
            return min(int(match.group(1)), 6)
        if candidate == "Title":
            return 1
    return None


def _extract_docx(path: Path, ocr_images: bool = False) -> str:
    """Extract text from a Word document (headings, paragraphs, tables).

    Headings are rendered as Markdown ``#`` lines and tables as
    pipe-separated rows, which keeps structure legible in the chunked
    text. python-docx exposes paragraphs and tables as two separate
    lists, so their relative order in the document is not preserved
    (tables come last).
    """
    try:
        import docx
    except ImportError as exc:
        raise ImportError(
            f"Reading .docx files requires the office extra: {_OFFICE_EXTRA_HINT}"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        level = _docx_heading_level(paragraph)
        parts.append(f"{'#' * level} {text}" if level else text)
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _extract_xlsx(path: Path, ocr_images: bool = False) -> str:
    """Extract cell values from an Excel workbook, one Markdown table per sheet.

    Formulas yield their last computed value (``data_only=True``); empty
    rows are dropped.
    """
    try:
        import openpyxl
    except ImportError as exc:
        raise ImportError(
            f"Reading .xlsx files requires the office extra: {_OFFICE_EXTRA_HINT}"
        ) from exc

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value).strip() for value in row]
                if any(cells):
                    rows.append("| " + " | ".join(cells) + " |")
            if rows:
                parts.append(f"## {sheet.title}\n\n" + "\n".join(rows))
    finally:
        workbook.close()
    return "\n\n".join(parts)


_LOADERS: dict[str, Callable[[Path, bool], str]] = {
    ".pdf": _extract_pdf,
    ".txt": _extract_plain_text,
    ".md": _extract_plain_text,
    ".markdown": _extract_plain_text,
    ".rst": _extract_plain_text,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".xlsm": _extract_xlsx,
}

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(_LOADERS)


def is_supported(path: str | Path) -> bool:
    """Return True if the file extension has a registered loader."""
    return Path(path).suffix.lower() in _LOADERS


def _resolve_loader(path: Path, table: dict[str, Callable[[Path, bool], str]]) -> Callable:
    if not path.is_file():
        raise FileNotFoundError(f"Document not found: {path}")
    loader = table.get(path.suffix.lower())
    if loader is None:
        raise ValueError(
            f"Unsupported file type {path.suffix!r} for {path.name}; "
            f"supported extensions: {', '.join(sorted(table))}"
        )
    return loader


def extract_text(path: str | Path, ocr_images: bool = False) -> str:
    """Extract text from a document.

    Args:
        path: Path to the source file.
        ocr_images: For PDFs, also OCR embedded images (best-effort).

    Returns:
        The extracted text (possibly empty for image-only documents
        when OCR is disabled).

    Raises:
        FileNotFoundError: If ``path`` does not exist.
        ValueError: If the file extension is not supported.
        ImportError: If the format needs an extra that is not installed.
    """
    path = Path(path)
    return _resolve_loader(path, _LOADERS)(path, ocr_images)
